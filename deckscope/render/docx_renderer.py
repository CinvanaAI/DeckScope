"""Word output — the format most people actually circulate."""
from __future__ import annotations

from pathlib import Path
from typing import Any, List

from .common import (ASSESSMENT_WORD, SEVERITY_WORD, as_list, findings_for,
                     header_block, txt)


def _hex(color: str) -> Any:
    from docx.shared import RGBColor
    return RGBColor.from_string(color.lstrip("#").upper())


def render(result, out_dir: Path, base: str, theme: str = "slate", **kw: Any) -> List[str]:
    try:
        import docx
        from docx.shared import Pt
    except ImportError:
        raise RuntimeError("Word output needs python-docx: pip install python-docx") from None

    from .common import theme as get_theme
    t = get_theme(theme)
    paths = []

    for lens, comp in result.comparisons.items():
        doc = docx.Document()
        h = header_block(result, lens)

        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(10.5)

        # Title block
        p = doc.add_paragraph()
        r = p.add_run(h["lens"].upper())
        r.font.size, r.font.bold, r.font.color.rgb = Pt(8.5), True, _hex(t["muted"])
        doc.add_heading(f"{result.company} — Deck vs. Market Analysis", level=0)

        # Same hierarchy as markdown and HTML: findings first, verdict demoted,
        # no composite score above the fold. See deckscope/findings.py.
        found = findings_for(result, lens)

        p = doc.add_paragraph()
        r = p.add_run(found.headline)
        r.font.size, r.font.bold, r.font.color.rgb = Pt(12), True, _hex(t["accent"])
        p = doc.add_paragraph()
        r = p.add_run(found.evidence_state)
        r.font.size, r.font.italic, r.font.color.rgb = Pt(9), True, _hex(t["muted"])

        if comp.get("integrity_note"):
            p = doc.add_paragraph()
            r = p.add_run(f"Integrity note. {comp['integrity_note']}")
            r.font.bold, r.font.color.rgb = True, _hex(t["bad"])

        tbl = doc.add_table(rows=0, cols=2)
        tbl.style = "Light Grid Accent 1"
        for k, v in (("Claims examined", str(found.counts.get("claims_examined", 0))),
                     ("Contested", str(found.counts.get("contested", 0))),
                     ("Omissions", str(found.counts.get("omissions", 0))),
                     ("Could not be checked",
                      str(found.counts.get("unverified", 0))),
                     ("Security screen",
                      (result.security or {}).get("overall_risk", "not run").upper()),
                     ("Sources", h["research"]), ("Model", h["model"]),
                     ("Generated", h["generated"])):
            row = tbl.add_row().cells
            row[0].text, row[1].text = str(k), str(v)
            row[0].paragraphs[0].runs[0].font.bold = True

        if found.contested:
            doc.add_heading("What the evidence contests", level=1)
            doc.add_paragraph(
                "Claims the deck makes that retrieved evidence pushes back on. "
                "Unsourced items are readings, not findings.")
            for f in found.contested:
                cites = (" ".join(f"[{s}]" for s in f.source_ids)
                         if f.source_ids else "(no source)")
                doc.add_paragraph(
                    f"{f.text} — {SEVERITY_WORD.get(f.severity, f.severity)}. "
                    f"{f.delta or f.why} {cites}", style="List Bullet")

        if found.omissions:
            doc.add_heading("What the deck leaves out", level=1)
            doc.add_paragraph(
                "Present in the market evidence, absent from the deck.")
            for f in found.omissions:
                doc.add_paragraph(f.text, style="List Bullet")

        if found.unverified:
            doc.add_heading("What could not be checked", level=1)
            doc.add_paragraph(
                "Neither confirmed nor refuted by the evidence retrieved. These "
                "are research tasks, not marks against the company.")
            for f in found.unverified:
                doc.add_paragraph(f.text, style="List Bullet")

        if found.next_steps:
            doc.add_heading("What to do next", level=1)
            for step in found.next_steps:
                doc.add_paragraph(step, style="List Number")

        # Summary
        doc.add_heading("Summary", level=1)
        for para in (comp.get("summary") or "").split("\n"):
            if para.strip():
                doc.add_paragraph(para.strip())

        # Verdict, demoted
        doc.add_heading("What this adds up to, for this lens", level=1)
        p = doc.add_paragraph()
        r = p.add_run(f"{h['verdict']} · confidence: {h['confidence']}")
        r.font.bold = True
        rationale = (comp.get("verdict") or {}).get("confidence_rationale")
        if rationale:
            p = doc.add_paragraph()
            r = p.add_run(f"Confidence basis: {rationale}")
            r.font.italic, r.font.size = True, Pt(9)
        p = doc.add_paragraph()
        r = p.add_run("A verdict is one reader's reading of the findings above, "
                      "through one lens. The findings are the durable part; this "
                      "line is not.")
        r.font.italic, r.font.size, r.font.color.rgb = True, Pt(9), _hex(t["muted"])

        # Scorecard
        rows = comp.get("scorecard") or []
        if rows:
            doc.add_heading("Scorecard", level=1)
            p = doc.add_paragraph()
            r = p.add_run(
                "Per-dimension, each with its reasoning. There is deliberately no "
                "headline total: a weighted average of seven subjective scores is "
                "the one figure here that cannot be traced to a source.")
            r.font.italic, r.font.size, r.font.color.rgb = (
                True, Pt(9), _hex(t["muted"]))
            tb = doc.add_table(rows=1, cols=4)
            tb.style = "Light Grid Accent 1"
            for i, head in enumerate(("Dimension", "Score", "Weight", "Why")):
                cell = tb.rows[0].cells[i]
                cell.text = head
                cell.paragraphs[0].runs[0].font.bold = True
            for r_ in rows:
                c = tb.add_row().cells
                c[0].text = txt(r_.get("dimension"))
                c[1].text = f"{txt(r_.get('score'))}/10"
                c[2].text = txt(r_.get("weight"))
                c[3].text = txt(r_.get("rationale"))

        # Claim audit
        audit = comp.get("claim_audit") or []
        if audit:
            doc.add_heading("Claim-by-claim audit", level=1)
            for c in audit:
                doc.add_heading(f"{txt(c.get('id'))} · {txt(c.get('claim'))}", level=2)
                a = (c.get("assessment") or "").lower()
                p = doc.add_paragraph()
                r = p.add_run(ASSESSMENT_WORD.get(a, a or "—"))
                r.font.bold = True
                r.font.color.rgb = _hex(
                    {"supported": t["good"], "partially-supported": t["warn"],
                     "contradicted": t["bad"]}.get(a, t["muted"]))
                for label, key in (("Market evidence", "market_evidence"),
                                   ("Gap", "delta"), ("So what", "so_what"),
                                   ("Evidence quality", "evidence_quality")):
                    if c.get(key):
                        p = doc.add_paragraph()
                        p.add_run(f"{label}: ").bold = True
                        p.add_run(str(c[key]))
                cites = _cites(c, result)
                p = doc.add_paragraph()
                p.add_run("Sources: ").bold = True
                p.add_run(cites)

        # Alignment
        align = comp.get("alignment") or {}
        if any(align.values()):
            doc.add_heading("Alignment", level=1)
            for key, title in (("where_deck_matches_market", "Deck matches the market"),
                               ("where_deck_overstates", "Deck overstates"),
                               ("where_deck_understates", "Deck understates"),
                               ("blind_spots", "Blind spots")):
                items = as_list(align.get(key))
                if items:
                    doc.add_heading(title, level=2)
                    for i in items:
                        doc.add_paragraph(str(i), style="List Bullet")

        # Risks
        risks = comp.get("risks") or []
        if risks:
            doc.add_heading("Risks", level=1)
            tb = doc.add_table(rows=1, cols=4)
            tb.style = "Light Grid Accent 1"
            for i, head in enumerate(("Risk", "Severity", "Likelihood", "Test")):
                cell = tb.rows[0].cells[i]
                cell.text = head
                cell.paragraphs[0].runs[0].font.bold = True
            for r_ in risks:
                c = tb.add_row().cells
                c[0].text = txt(r_.get("risk"))
                c[1].text = txt(r_.get("severity"))
                c[2].text = txt(r_.get("likelihood"))
                c[3].text = txt(r_.get("mitigation_or_test"))

        # Questions and actions already appear, consolidated and ranked, under
        # "What to do next" near the top. Only the owner/priority detail that
        # the actions table alone carried is repeated here.
        acts = comp.get("actions") or []
        if acts:
            doc.add_heading("Who does what", level=1)
            tb = doc.add_table(rows=1, cols=3)
            tb.style = "Light Grid Accent 1"
            for i, head in enumerate(("Priority", "Action", "Owner")):
                cell = tb.rows[0].cells[i]
                cell.text = head
                cell.paragraphs[0].runs[0].font.bold = True
            for a in sorted(acts, key=lambda x: str(x.get("priority", "P9"))):
                c = tb.add_row().cells
                c[0].text = txt(a.get("priority"))
                c[1].text = txt(a.get("action"))
                c[2].text = txt(a.get("owner"))

        # ---- The analysis sections added after the original report design.
        # Word is the format people circulate, so it carries the decision-relevant
        # ones. PPTX deliberately stays a curated summary; that is documented.
        _docx_discovery(doc, result, t)
        _docx_market_structure(doc, result.market, t)
        _docx_opportunity(doc, result, t)

        # References — every source, cited or not
        doc.add_page_break()
        doc.add_heading("References", level=1)
        reg = getattr(result, "registry", None)
        if not reg or not reg.sources:
            doc.add_paragraph(
                "No external sources were retrieved for this analysis. Every statement "
                "above rests on the model's training knowledge and on the deck itself, "
                "and should be treated as unverified.")
        else:
            st = reg.stats()
            doc.add_paragraph(
                f"{st['total']} sources retrieved and screened: {st['cited']} cited, "
                f"{st['consulted_uncited']} consulted without being cited, "
                f"{st['quarantined']} dropped by the security screen. All are listed "
                f"below so that absence of evidence is as visible as its presence.")
            for group, title in ((reg.cited, "Cited in this analysis"),
                                 (reg.consulted, "Consulted, not cited"),
                                 (reg.quarantined, "Dropped by the security screen")):
                if not group:
                    continue
                doc.add_heading(title, level=2)
                for s in group:
                    p = doc.add_paragraph(style="List Number")
                    p.add_run(f"[{s.sid}] ").bold = True
                    p.add_run(s.title or s.url or "(untitled)")
                    if s.url:
                        r = p.add_run(f" — {s.url}")
                        r.font.size, r.font.color.rgb = Pt(8.5), _hex(t["muted"])
                    tail = " · ".join(x for x in (s.published, s.reliability,
                                                  "; ".join(s.cited_by[:3]), s.note) if x)
                    if tail:
                        r = p.add_run(f"\n{tail}")
                        r.font.size, r.font.color.rgb = Pt(8.5), _hex(t["muted"])

        # Security
        sec = result.security or {}
        if sec:
            doc.add_heading("Input integrity screen", level=1)
            p = doc.add_paragraph()
            r = p.add_run(f"Overall risk: {sec.get('overall_risk', 'clean').upper()} "
                          f"(mode: {sec.get('mode', 'balanced')})")
            r.font.bold = True
            if sec.get("overall_risk") == "clean":
                doc.add_paragraph(
                    "The deck and every web source were screened for content written to "
                    "influence the AI rather than inform a human reader. Nothing was found.")
            else:
                for key, title in (("deck", "Pitch deck"), ("web_sources", "Web sources")):
                    findings = (sec.get(key) or {}).get("findings") or []
                    if not findings:
                        continue
                    doc.add_heading(title, level=2)
                    for f in findings[:30]:
                        doc.add_paragraph(
                            f"[{f.get('severity')}] {f.get('where')} — {f.get('detail')} "
                            f"({f.get('action')})", style="List Bullet")

        p = doc.add_paragraph()
        r = p.add_run(f"Generated by DeckScope · "
                      f"{h['model']} · {h['research']}. AI-generated analysis: verify every "
                      f"figure before relying on it. Not investment advice.")
        r.font.size, r.font.italic, r.font.color.rgb = Pt(8), True, _hex(t["muted"])

        path = out_dir / f"{base}_{lens}.docx"
        doc.save(str(path))
        paths.append(str(path))
    return paths


def _docx_discovery(doc: Any, result: Any, t: Any) -> None:
    """What the cold pass found that the claim-directed pass never looked for."""
    delta = getattr(result, "discovery_delta", None) or {}
    if not delta or not delta.get("ran"):
        return
    doc.add_heading("What the deck steered the research away from", level=1)
    doc.add_paragraph(
        "The market analysis was given the deck's claims — it has to be, it is "
        "checking them — so its searches were shaped by what the deck raises. The "
        "category was therefore also researched cold, by a pass that saw only the "
        "category name and never a single claim.")
    if not delta.get("anything_found"):
        doc.add_paragraph(
            "Nothing. Researching the category from scratch surfaced no competitor, "
            "headwind or absorber the claim-directed pass had missed.")
        return
    doc.add_paragraph(
        f"Overlap between the two routes: "
        f"{float(delta.get('agreement') or 0):.0%} of the competitors named.")

    only_cold = delta.get("competitors_only_cold") or []
    if only_cold:
        doc.add_heading("Competitors found only when the deck was out of the room",
                        level=2)
        table = doc.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 1"
        for i, head in enumerate(("Competitor", "Position", "Threat")):
            cell = table.rows[0].cells[i]
            cell.text = head
            cell.paragraphs[0].runs[0].font.bold = True
        for c in only_cold:
            row = table.add_row().cells
            row[0].text = txt(c.get("name"))
            row[1].text = txt(c.get("position"))
            row[2].text = txt(c.get("threat_level"))

    for key, label in (("absorbers_only_cold", "Potential absorbers not raised"),
                       ("adjacent_only_cold", "Adjacent markets missed")):
        items = delta.get(key) or []
        if items:
            para = doc.add_paragraph()
            para.add_run(f"{label}: ").bold = True
            para.add_run(", ".join(str(x) for x in items))

    for h in (delta.get("headwinds_only_cold") or []):
        doc.add_paragraph(str(h), style="List Bullet")


def _docx_market_structure(doc: Any, market: Any, t: Any) -> None:
    """Saturation, absorption risk and the open-source signal."""
    land = (market or {}).get("competitive_landscape") or {}
    sat = land.get("saturation") or {}
    ab = (market or {}).get("absorption_risk") or {}
    assessment = (market or {}).get("bundling_assessment") or {}
    if not any([any(sat.values()), any(ab.values()), assessment]):
        return

    doc.add_heading("Market structure", level=1)

    if any(sat.values()):
        doc.add_heading("Saturation", level=2)
        table = doc.add_table(rows=0, cols=2)
        table.style = "Light Grid Accent 1"
        for label, key in (("Funded competitors", "funded_competitors_known"),
                           ("New entrants", "new_entrants_trend"),
                           ("Pricing", "pricing_direction"),
                           ("Consolidation", "consolidation_activity"),
                           ("Lifecycle stage", "lifecycle_stage"),
                           ("Room to enter", "room_for_a_new_entrant")):
            row = table.add_row().cells
            row[0].text = label
            row[1].text = txt(sat.get(key))
            row[0].paragraphs[0].runs[0].font.bold = True
        if sat.get("why"):
            doc.add_paragraph(str(sat["why"]))

    if any(ab.values()):
        doc.add_heading("Is this a product or a feature?", level=2)
        para = doc.add_paragraph()
        run = para.add_run(f"{txt(ab.get('verdict')).upper()} · horizon "
                           f"{txt(ab.get('horizon'))} · confidence "
                           f"{txt(ab.get('confidence'))}")
        run.font.bold = True
        doc.add_paragraph(
            "Categories are regularly built out by startups, proven useful, and then "
            "bundled into a platform that already owns the customer. When that "
            "happens the market stops existing separately.")
        for a in (ab.get("likely_absorbers") or []):
            para = doc.add_paragraph(style="List Bullet")
            para.add_run(f"{txt(a.get('name'))} — {txt(a.get('mechanism'))}. ").bold = True
            para.add_run(txt(a.get("why_them")))
        for x in as_list(ab.get("what_would_prevent_it")):
            doc.add_paragraph(f"Would prevent it: {x}", style="List Bullet")

    if assessment and assessment.get("applicable") is not False:
        doc.add_heading("Open source, and what it predicts", level=2)
        para = doc.add_paragraph()
        para.add_run(f"Bundling risk from commoditization: "
                     f"{txt(assessment.get('level')).upper()}").bold = True
        if assessment.get("reasoning"):
            doc.add_paragraph(str(assessment["reasoning"]))
        for c in (assessment.get("caveats") or []):
            doc.add_paragraph(str(c), style="List Bullet")


def _docx_opportunity(doc: Any, result: Any, t: Any) -> None:
    """What buying the listed alternative would require instead."""
    from docx.shared import Pt

    opp = getattr(result, "opportunity", None) or {}
    if not opp or opp.get("error"):
        return

    doc.add_heading("Compared to what?", level=1)
    if opp.get("headline"):
        para = doc.add_paragraph()
        run = para.add_run(str(opp["headline"]))
        run.font.italic = True
        run.font.color.rgb = _hex(t["accent"])
    doc.add_paragraph(
        "This is not a forecast. It is the outcome this company would have to reach "
        "to match each benchmark, under stated assumptions you can change.")

    comps = opp.get("comparables") or []
    if comps:
        doc.add_heading("The named competitors, and whether you could buy them instead",
                        level=2)
        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Grid Accent 1"
        for i, head in enumerate(("Competitor", "Listed", "Market cap",
                                  "5-year return")):
            cell = table.rows[0].cells[i]
            cell.text = head
            cell.paragraphs[0].runs[0].font.bold = True
        for c in comps:
            row = table.add_row().cells
            row[0].text = txt(c.get("name"))
            row[1].text = txt(c.get("ticker")) if c.get("ticker") else "private"
            row[2].text = txt(c.get("market_cap_display"))
            row[3].text = (f"{c['total_return_5y']}x"
                           if c.get("total_return_5y") else "—")

    reqs = opp.get("requirements") or {}
    if reqs:
        doc.add_heading("What this company would have to reach", level=2)
        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Grid Accent 1"
        for i, head in enumerate(("To match", "Exit value needed", "Implied revenue",
                                  "Multiple of today")):
            cell = table.rows[0].cells[i]
            cell.text = head
            cell.paragraphs[0].runs[0].font.bold = True
        for label, r in reqs.items():
            row = table.add_row().cells
            row[0].text = str(label)
            row[1].text = txt(r.get("exit_value_required_display"))
            row[2].text = txt(r.get("implied_arr_required_display"))
            row[3].text = f"{txt(r.get('growth_multiple_required'))}x"

    a = opp.get("assumptions") or {}
    if a:
        para = doc.add_paragraph()
        para.add_run("Assumptions: ").bold = True
        para.add_run(
            f"{float(a.get('future_dilution') or 0):.0%} future dilution · "
            f"{txt(a.get('exit_revenue_multiple'))}x exit revenue multiple · "
            f"{txt(a.get('horizon_years'))} years · "
            f"{txt(a.get('preference_stack'))}x preference. "
            f"Change any of these and every figure changes.")

    for r in (opp.get("base_rates") or []):
        doc.add_paragraph(
            f"{txt(r.get('statement'))}: {txt(r.get('value'))} "
            f"({txt(r.get('population'))}) — {txt(r.get('caveat'))}",
            style="List Bullet")

    para = doc.add_paragraph()
    run = para.add_run(str(opp.get("disclaimer") or ""))
    run.font.size, run.font.italic = Pt(8.5), True
    run.font.color.rgb = _hex(t["muted"])


def _cites(claim: Any, result) -> str:
    reg = getattr(result, "registry", None)
    parts, seen = [], set()
    for ref in list(as_list(claim.get("source_ids"))) + list(as_list(claim.get("sources"))):
        if not ref:
            continue
        s = reg.find(str(ref)) if reg else None
        if s and s.sid not in seen:
            seen.add(s.sid)
            parts.append(f"[{s.sid}] {s.url or s.title}")
        elif not s:
            parts.append(str(ref))
    return "; ".join(parts) if parts else "none cited — this assessment rests on no source"
